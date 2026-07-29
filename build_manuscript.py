"""Generate a professional research manuscript (.docx) for the Yashi PROTAC project."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "outputs" / "figures"
OUT = ROOT / "Manuscript_PROTAC_Substructure_Segmentation.docx"


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, space_before=0,
             first_line=True):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading_custom(doc, text, level=1):
    """Numbered-style section headings in Times New Roman."""
    sizes = {1: 14, 2: 12, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True,
                 italic=(level >= 3))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True)
        set_cell_border(hdr[i])

    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            set_cell_border(cells[c_i])

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_figure(doc, path: Path, width_in=5.8, caption=""):
    if not path.exists():
        add_para(doc, f"[Figure missing: {path.name}]", italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if caption:
        add_caption(doc, caption)


def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # =====================================================================
    # TITLE PAGE
    # =====================================================================
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Graph Neural Network-Based Substructure Segmentation of "
        "PROTAC Molecules with Weak Supervision and Reconstruction-Aware Learning"
    )
    set_run_font(run, size=16, bold=True)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.5

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "A Computational Study on Atom-Level Warhead–Linker–E3 Ligand Assignment "
        "Using PROTAC-DB 4.0"
    )
    set_run_font(run, size=12, italic=True)
    p.paragraph_format.space_after = Pt(24)

    add_para(doc, "Author Name(s)", align=WD_ALIGN_PARAGRAPH.CENTER,
             first_line=False, bold=True)
    add_para(doc, "Affiliation / Laboratory / Institution",
             align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, italic=True)
    add_para(doc, "Email: corresponding.author@institution.edu",
             align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "Date: July 2026", align=WD_ALIGN_PARAGRAPH.CENTER,
             first_line=False, space_after=24)

    add_para(
        doc,
        "This manuscript presents an original computational methodology for "
        "automatic identification of warhead, linker, and E3 ligand substructures "
        "within proteolysis-targeting chimera (PROTAC) molecules. The work "
        "integrates weak supervision, graph neural networks, reconstruction-aware "
        "training objectives, and out-of-distribution evaluation protocols.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        italic=True,
        first_line=False,
        space_after=24,
    )

    doc.add_page_break()

    # =====================================================================
    # ABSTRACT
    # =====================================================================
    add_heading_custom(doc, "Abstract", level=1)
    add_para(
        doc,
        "Proteolysis-targeting chimeras (PROTACs) are heterobifunctional molecules "
        "that recruit an E3 ubiquitin ligase to a protein of interest, thereby "
        "inducing targeted protein degradation. Rational design and retrospective "
        "analysis of PROTACs require reliable identification of the three chemical "
        "substructures that constitute each molecule: the warhead, the linker, and "
        "the E3 ligand. Manual annotation does not scale to large databases, while "
        "existing automated strategies either operate at the bond-cut or "
        "sequence-to-sequence level and typically evaluate only random held-out "
        "splits. In this work, we formulate PROTAC substructure identification as "
        "atom-level three-class segmentation on molecular graphs. Weak labels are "
        "generated by dictionary-based substructure matching and retained only when "
        "the implied three-way fragmentation reassembles to the original molecule. "
        "A Graph Isomorphism Network (GIN) is trained with a multi-task objective "
        "combining atom classification, bond-boundary prediction, fragment "
        "smoothness, and a differentiable soft-reconstruction loss. We further "
        "compare training from scratch with AttrMask-style self-supervised "
        "pretraining of the same encoder. Evaluation is performed on five splits, "
        "including unseen-warhead, unseen-E3, fingerprint-dissimilar, and "
        "newer-chemotype regimes, against dictionary and XGBoost baselines. On "
        "random test molecules, the scratch GNN achieves 0.922 atom accuracy and "
        "0.917 reassembly accuracy. Soft reconstruction training raises unseen-E3 "
        "reassembly from approximately 0.33 in an earlier configuration to 0.81. "
        "Across all splits, the GNN substantially outperforms the XGBoost "
        "bond-cutting baseline on atom-level accuracy. The results support "
        "reconstruction-aware graph learning as a practical and scientifically "
        "grounded route to PROTAC substructure segmentation.",
        first_line=False,
    )
    add_para(
        doc,
        "Keywords: PROTAC; targeted protein degradation; graph neural network; "
        "weak supervision; substructure segmentation; molecular reconstruction; "
        "out-of-distribution evaluation",
        italic=True,
        first_line=False,
        space_before=8,
    )

    # =====================================================================
    # 1. INTRODUCTION
    # =====================================================================
    add_heading_custom(doc, "1. Introduction", level=1)

    add_para(
        doc,
        "Targeted protein degradation has emerged as a transformative modality in "
        "chemical biology and drug discovery. Unlike classical occupancy-driven "
        "inhibitors, degradation-based therapeutics remove disease-relevant "
        "proteins from the cell, often enabling sustained pharmacological effects "
        "at lower occupancy and expanding the accessible target space to proteins "
        "that are difficult to inhibit conventionally [1,2]. Among degrader "
        "architectures, proteolysis-targeting chimeras (PROTACs) are the most "
        "widely studied class of heterobifunctional degraders [3,4].",
    )
    add_para(
        doc,
        "A canonical PROTAC comprises three covalently connected components: "
        "(i) a warhead (also called the protein-of-interest ligand) that binds the "
        "target protein; (ii) an E3 ligand that engages an E3 ubiquitin ligase such "
        "as cereblon (CRBN), von Hippel–Lindau (VHL), or inhibitors of apoptosis "
        "proteins (IAP); and (iii) a linker that bridges the two ligands and "
        "strongly influences ternary-complex geometry, permeability, and "
        "degradation efficiency [5,6]. Because linker and ligand identity jointly "
        "determine biological performance, automated extraction of these "
        "substructures from large PROTAC collections is essential for "
        "structure–activity learning, generative design, and database curation.",
    )
    add_para(
        doc,
        "Despite the rapid growth of public PROTAC resources such as PROTAC-DB "
        "[7], atom-level annotations of warhead, linker, and E3 ligand regions "
        "remain scarce. Chemist-drawn gold labels are expensive to produce at "
        "database scale. Consequently, prior computational work has relied on "
        "dictionary matching, bond-level machine learning, or sequence-to-sequence "
        "models trained on curated or synthetically expanded splitting datasets "
        "[8,9]. These approaches provide valuable baselines, yet several gaps "
        "remain for a segmentation-oriented research agenda: (i) limited use of "
        "atom-level multi-task graph learning; (ii) reconstruction treated mainly "
        "as post hoc validation rather than a training objective; and "
        "(iii) insufficient out-of-distribution (OOD) testing beyond random splits.",
    )
    add_para(
        doc,
        "This manuscript addresses those gaps. We present a complete methodology "
        "for PROTAC substructure segmentation that (1) constructs reassembly-filtered "
        "weak labels, (2) trains a compact GIN with an explicit soft-reconstruction "
        "loss, (3) compares scratch versus AttrMask-pretrained encoders, and "
        "(4) evaluates performance under five complementary data partitions, "
        "including chemotype-held-out and fingerprint-dissimilar regimes. The "
        "central scientific claim is that atom-level graph segmentation with "
        "reconstruction-aware supervision yields more transferable chemical "
        "understanding than bond-cut classification alone, while remaining "
        "reproducible on modest computational resources.",
    )

    add_heading_custom(doc, "1.1 Research problem", level=2)
    add_para(
        doc,
        "Formally, let a PROTAC be represented by a molecular graph G = (V, E), "
        "where each node v ∈ V is an atom and each edge e ∈ E is a chemical bond. "
        "The task is to assign every atom a label y_v ∈ {warhead, linker, E3 ligand} "
        "such that the predicted partition induces a chemically meaningful "
        "three-fragment decomposition of the original molecule. The research "
        "problem therefore combines supervised learning under noisy pseudo-labels "
        "with chemistry-aware constraints on connectivity and reassembly.",
    )

    add_heading_custom(doc, "1.2 Contributions", level=2)
    add_para(
        doc,
        "The main contributions of this study are as follows. First, we cast "
        "PROTAC substructure identification as atom-level segmentation and release "
        "a compact, fully reproducible pipeline on PROTAC-DB 4.0. Second, we "
        "introduce a soft-reconstruction training objective that encourages "
        "canonical linker attachment topology and simultaneous presence of all "
        "three fragment classes. Third, we establish an OOD evaluation suite "
        "comprising unseen warhead, unseen E3, fingerprint-dissimilar, and "
        "newer-chemotype splits. Fourth, we provide a controlled comparison "
        "between a randomly initialized GIN and an AttrMask-pretrained GIN against "
        "dictionary and XGBoost baselines under identical metrics.",
        first_line=True,
    )

    # =====================================================================
    # 2. RELATED WORK
    # =====================================================================
    add_heading_custom(doc, "2. Related Work", level=1)

    add_heading_custom(doc, "2.1 PROTAC chemistry and databases", level=2)
    add_para(
        doc,
        "PROTACs operate by inducing proximity between a target protein and an E3 "
        "ligase, leading to ubiquitination and proteasomal degradation [1–4]. "
        "Design principles emphasize ligand affinity, linker length and "
        "composition, ternary-complex cooperativity, and physicochemical "
        "properties related to cellular uptake [5,6]. PROTAC-DB aggregates "
        "structures, activities, and component libraries for warheads, linkers, "
        "and E3 ligands, and therefore provides a natural substrate for "
        "data-driven substructure analysis [7].",
    )

    add_heading_custom(doc, "2.2 Automated PROTAC splitting", level=2)
    add_para(
        doc,
        "Ribes and colleagues recently introduced PROTAC-Splitter, a machine "
        "learning framework that combines dictionary-based curation with learned "
        "models for automated identification of PROTAC substructures, including "
        "transformer-based sequence models and classical bond classifiers [8]. "
        "Complementary laboratory work has explored XGBoost-based prediction of "
        "linker-cutting bonds using engineered topological and chemical bond "
        "features [9]. These studies demonstrate that automated splitting is "
        "feasible, but they primarily emphasize bond or string outputs and "
        "random-split evaluation. Our formulation instead predicts an atom-wise "
        "segmentation map and stresses reconstruction-aware learning together "
        "with explicit OOD partitions.",
    )

    add_heading_custom(doc, "2.3 Graph neural networks for molecules", level=2)
    add_para(
        doc,
        "Graph neural networks have become a standard representation learning "
        "tool for molecules because they operate directly on atoms and bonds "
        "[10,11]. The Graph Isomorphism Network (GIN) is particularly expressive "
        "among message-passing architectures [12]. Self-supervised molecular "
        "pretraining strategies such as attribute masking (AttrMask) and later "
        "refinements including Mole-BERT aim to improve transfer to downstream "
        "tasks [13,14]. In this study we adopt a lightweight AttrMask pretraining "
        "protocol on the same GIN backbone used for segmentation, enabling a "
        "fair scratch-versus-pretrained comparison without requiring specialized "
        "GPU-only toolchains.",
    )

    # =====================================================================
    # 3. MATERIALS AND METHODS
    # =====================================================================
    add_heading_custom(doc, "3. Materials and Methods", level=1)

    add_heading_custom(doc, "3.1 Dataset", level=2)
    add_para(
        doc,
        "All experiments use PROTAC-DB 4.0 files comprising PROTAC structures "
        "together with warhead and E3 ligand reference libraries [7]. Molecules "
        "were parsed with RDKit [15]. PROTACs outside the range of 15–120 heavy "
        "atoms were discarded to exclude trivial fragments and extremely large "
        "outliers. For computational tractability in the reported runs, a "
        "reproducible random sample of 6,000 candidate SMILES strings was "
        "processed for weak labeling; after reassembly filtering, 1,674 labeled "
        "PROTACs remained for learning and evaluation.",
    )

    add_heading_custom(doc, "3.2 Weak label generation with reassembly filtering", level=2)
    add_para(
        doc,
        "Because gold atom-level annotations are unavailable at scale, we "
        "construct weak labels by substructure matching. For each PROTAC, the "
        "largest non-overlapping matches from the warhead and E3 libraries are "
        "selected; unmatched atoms are assigned to the linker class. A candidate "
        "labeling is accepted only if cutting every inter-class bond yields "
        "exactly three chemically sanitizable fragments whose atom counts are "
        "consistent with reassembly of the parent molecule. This filter follows "
        "the scientific intuition of curation-by-reconstruction used in recent "
        "PROTAC splitting work [8], while producing atom-level supervision for "
        "segmentation.",
    )

    add_heading_custom(doc, "3.3 Molecular graph representation", level=2)
    add_para(
        doc,
        "Each labeled PROTAC is converted to an undirected molecular graph. Node "
        "features encode atom type, hybridization, degree, formal charge, "
        "implicit hydrogens, aromaticity, and ring membership. Edge features "
        "encode bond type, conjugation, and ring membership. Graphs are batched "
        "by block-diagonal concatenation for mini-batch training without "
        "dependence on PyTorch Geometric.",
    )

    add_heading_custom(doc, "3.4 Model architecture", level=2)
    add_para(
        doc,
        "The segmentation model is a four-layer GIN encoder with hidden width 64, "
        "followed by three heads: (i) an atom classifier over {warhead, linker, "
        "E3}, (ii) a bond-boundary classifier, and (iii) an AttrMask head used "
        "only during self-supervised pretraining. The architecture contains on "
        "the order of tens of thousands of parameters and trains comfortably on "
        "CPU, supporting reproducible experimentation in typical laboratory "
        "settings.",
    )

    add_heading_custom(doc, "3.5 Training objectives", level=2)
    add_para(
        doc,
        "The supervised fine-tuning loss is a weighted sum of four terms:",
        first_line=True,
    )
    add_para(
        doc,
        "L = L_atom + 0.5 L_bond + 0.1 L_smooth + 0.1 L_recon,",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        first_line=False,
        space_before=6,
        space_after=6,
    )
    add_para(
        doc,
        "where L_atom is class-weighted cross-entropy over atom labels (with an "
        "additional boost on the minority linker class), L_bond is cross-entropy "
        "over whether a bond connects different atom classes, L_smooth penalizes "
        "disagreement between neighboring atoms that share the same weak label, "
        "and L_recon is a soft-reconstruction objective. The reconstruction term "
        "combines (a) a linker-attachment proxy encouraging approximately two "
        "linker–nonlinker boundary contacts per molecule and (b) a class-presence "
        "proxy encouraging nonzero soft mass for all three classes. Together, "
        "these differentiable surrogates push the network toward partitions that "
        "can reassemble, rather than optimizing atom accuracy alone.",
        first_line=False,
    )

    add_heading_custom(doc, "3.6 AttrMask pretraining", level=2)
    add_para(
        doc,
        "For the pretrained variant, the shared encoder is first optimized with "
        "an AttrMask objective [13]: a random subset of atom feature vectors is "
        "zeroed, and the network predicts atom type from the remaining graph "
        "context. After five pretraining epochs on all available graphs, the "
        "encoder is fine-tuned with the segmentation objective above. A matched "
        "scratch model is trained identically except for random initialization. "
        "This design isolates the effect of pretraining while remaining faithful "
        "to the Mole-BERT-family philosophy of masked atom modeling [14], without "
        "requiring external pretrained checkpoints.",
    )

    add_heading_custom(doc, "3.7 Post-hoc chemistry-constrained decoding", level=2)
    add_para(
        doc,
        "At inference, raw argmax predictions may contain small disconnected "
        "islands of a given class. We therefore apply a lightweight constraint: "
        "retain the largest connected component of each predicted class and "
        "reassign remaining atoms to the nearest retained neighbor class. All "
        "primary tables report constrained metrics; raw metrics are retained for "
        "transparency.",
    )

    add_heading_custom(doc, "3.8 Baselines", level=2)
    add_para(
        doc,
        "Two baselines are evaluated under identical splits and metrics. The "
        "dictionary baseline re-applies library matching at test time and is "
        "reported as a ceiling reference; because the same style of matching "
        "generated the weak labels, near-perfect dictionary scores are expected "
        "and must be interpreted as circular rather than as a fair competitor. "
        "The XGBoost baseline follows the bond-cutting paradigm [8,9]: bonds are "
        "scored with topological and chemical descriptors, the top cuts are "
        "selected to produce three fragments, and fragments are mapped to "
        "warhead/linker/E3 labels for metric computation.",
    )

    # =====================================================================
    # 4. EXPERIMENTAL SETUP
    # =====================================================================
    add_heading_custom(doc, "4. Experimental Setup", level=1)

    add_heading_custom(doc, "4.1 Data splits", level=2)
    add_para(
        doc,
        "Five complementary partitions are used. (1) Random: an 80/10/10 "
        "shuffle split measuring in-distribution performance. (2) Unseen "
        "warhead: rare warhead references are held out for validation and "
        "test. (3) Unseen E3: rare E3 references are held out analogously. "
        "(4) Fingerprint OOD: molecules with low maximum Morgan fingerprint "
        "Tanimoto similarity to a training pool are reserved for test, "
        "approximating novel chemical sources. (5) Newer chemotype: the "
        "highest PROTAC-DB Compound IDs form the test set, approximating "
        "newer database entries. Approximate test sizes are 167–184 molecules "
        "per split.",
    )

    add_heading_custom(doc, "4.2 Optimization and implementation", level=2)
    add_para(
        doc,
        "Models were implemented in PyTorch [16] with RDKit chemistry support "
        "[15]. Fine-tuning used Adam with learning rate 2×10⁻³, batch size 32, "
        "gradient clipping at norm 5, and early selection by validation "
        "reassembly plus a small atom-accuracy term. AttrMask pretraining used "
        "learning rate 1×10⁻³ for five epochs. All reported runs were executed "
        "on CPU to emphasize accessibility and reproducibility.",
    )

    add_heading_custom(doc, "4.3 Evaluation metrics", level=2)
    add_para(
        doc,
        "We report three complementary metrics. Atom accuracy is the fraction of "
        "atoms whose predicted class matches the weak label. Exact-3-fragment "
        "rate is the fraction of molecules whose predicted boundary cuts produce "
        "exactly three fragments. Reassembly accuracy is the fraction of "
        "molecules for which those fragments are chemically valid and consistent "
        "with reconstructing the parent structure. Reassembly is treated as the "
        "primary chemistry-facing criterion.",
    )

    add_figure(
        doc,
        FIG / "04_atom_acc_bars.png",
        width_in=5.9,
        caption="Figure 1. Atom accuracy of dictionary, XGBoost, scratch GNN, and "
                "AttrMask-pretrained GNN across five evaluation splits.",
    )
    add_figure(
        doc,
        FIG / "04_reassembly_bars.png",
        width_in=5.9,
        caption="Figure 2. Reassembly accuracy across the same models and splits. "
                "Dictionary scores are near-perfect by construction and should be "
                "interpreted as a circular ceiling.",
    )

    # =====================================================================
    # 5. RESULTS
    # =====================================================================
    add_heading_custom(doc, "5. Results", level=1)

    add_heading_custom(doc, "5.1 Overall comparison", level=2)
    add_para(
        doc,
        "Table 1 summarizes constrained test metrics for all models and splits. "
        "On the random split, the scratch GNN attains 0.922 atom accuracy, 0.988 "
        "exact-3 rate, and 0.917 reassembly accuracy, exceeding XGBoost "
        "(0.765 / 0.833 / 0.833). The AttrMask-pretrained GNN is competitive "
        "(0.915 / 0.982 / 0.875) but not uniformly superior to scratch training "
        "on this in-distribution partition.",
    )

    add_caption(
        doc,
        "Table 1. Test-set performance under chemistry-constrained decoding. "
        "Dictionary results are circular with respect to weak-label generation.",
    )
    add_table(
        doc,
        ["Split", "Model", "Atom acc.", "Exact-3", "Reassembly"],
        [
            ["Random", "Dictionary*", "0.976", "1.000", "1.000"],
            ["Random", "XGBoost", "0.765", "0.833", "0.833"],
            ["Random", "GNN scratch", "0.922", "0.988", "0.917"],
            ["Random", "GNN pretrained", "0.915", "0.982", "0.875"],
            ["Unseen warhead", "XGBoost", "0.649", "0.895", "0.895"],
            ["Unseen warhead", "GNN scratch", "0.901", "0.988", "0.772"],
            ["Unseen warhead", "GNN pretrained", "0.874", "1.000", "0.912"],
            ["Unseen E3", "XGBoost", "0.640", "0.918", "0.918"],
            ["Unseen E3", "GNN scratch", "0.887", "0.967", "0.810"],
            ["Unseen E3", "GNN pretrained", "0.853", "0.951", "0.674"],
            ["Fingerprint OOD", "XGBoost", "0.650", "0.856", "0.856"],
            ["Fingerprint OOD", "GNN scratch", "0.853", "0.880", "0.731"],
            ["Fingerprint OOD", "GNN pretrained", "0.866", "0.976", "0.689"],
            ["Newer chemotype", "XGBoost", "0.532", "0.617", "0.617"],
            ["Newer chemotype", "GNN scratch", "0.923", "0.898", "0.784"],
            ["Newer chemotype", "GNN pretrained", "0.920", "0.976", "0.844"],
        ],
    )
    add_para(
        doc,
        "*Dictionary performance is inflated because weak labels were produced "
        "by the same matching paradigm.",
        italic=True,
        first_line=False,
        size=10,
        space_before=0,
        space_after=10,
    )

    add_heading_custom(doc, "5.2 Generalization under chemotype shift", level=2)
    add_para(
        doc,
        "Under unseen-warhead partitioning, atom accuracy remains high for both "
        "GNN variants (0.901 scratch; 0.874 pretrained), whereas XGBoost drops "
        "to 0.649. Pretraining yields the best reassembly in this regime "
        "(0.912), suggesting that masked-atom initialization can help recover "
        "coherent fragment topology when warhead scaffolds are novel. On "
        "unseen-E3 molecules, scratch training is strongest (0.887 atom "
        "accuracy; 0.810 reassembly). Relative to an earlier experimental "
        "configuration without the full soft-reconstruction objective, "
        "unseen-E3 reassembly improved from approximately 0.33 to 0.81, "
        "indicating that reconstruction-aware supervision is decisive for "
        "difficult E3 chemotypes.",
    )

    add_heading_custom(doc, "5.3 Fingerprint OOD and newer chemotypes", level=2)
    add_para(
        doc,
        "Fingerprint-dissimilar molecules constitute a deliberately hard test. "
        "Both GNN variants retain atom accuracies above 0.85, outperforming "
        "XGBoost (0.650). Reassembly is moderate (0.731 scratch; 0.689 "
        "pretrained), confirming that chemical novelty remains challenging even "
        "for graph models. On newer Compound-ID holdings, XGBoost collapses to "
        "0.532 atom accuracy, while GNN scratch and pretrained models remain "
        "near 0.92. Pretraining again improves reassembly (0.844 versus 0.784), "
        "consistent with a useful inductive bias for newer database entries.",
    )

    add_figure(
        doc,
        FIG / "03_scratch_vs_pretrained.png",
        width_in=5.8,
        caption="Figure 3. Validation reassembly trajectories for scratch versus "
                "AttrMask-pretrained GNN on selected splits.",
    )

    add_heading_custom(doc, "5.4 Qualitative segmentation examples", level=2)
    add_para(
        doc,
        "Figures 4 and 5 illustrate constrained predictions on held-out PROTACs. "
        "Warhead, linker, and E3 regions appear as contiguous colored overlays, "
        "supporting the interpretation that the model learns spatially coherent "
        "roles rather than scattered atom tags. Examples drawn from "
        "fingerprint-OOD and unseen-warhead tests further indicate that the "
        "segmentation remains chemically readable under distribution shift.",
    )
    add_figure(
        doc,
        FIG / "05_demo_0.png",
        width_in=5.2,
        caption="Figure 4. Example constrained segmentation on a random-split "
                "test PROTAC (red = warhead, blue = linker, green = E3 ligand).",
    )
    demo_ood = FIG / "05_demo_fp_ood.png"
    if not demo_ood.exists():
        demo_ood = FIG / "05_demo_unseen_wh.png"
    add_figure(
        doc,
        demo_ood,
        width_in=5.2,
        caption="Figure 5. Example segmentation under an out-of-distribution "
                "test condition.",
    )

    # =====================================================================
    # 6. DISCUSSION
    # =====================================================================
    add_heading_custom(doc, "6. Discussion", level=1)

    add_heading_custom(doc, "6.1 Comparison with prior automated approaches", level=2)
    add_para(
        doc,
        "Relative to XGBoost bond-cutting methods in the spirit of laboratory "
        "and PROTAC-Splitter baselines [8,9], the proposed GNN is consistently "
        "stronger on atom-level accuracy and more stable under newer-chemotype "
        "and unseen-warhead shifts. XGBoost can retain high exact-3 and "
        "reassembly numbers because it is constrained to select a small number "
        "of cuts; however, those metrics can remain high even when atom-role "
        "assignment is poor. Atom accuracy therefore provides a stricter view "
        "of segmentation quality.",
    )
    add_para(
        doc,
        "Direct numerical comparison with transformer-based PROTAC-Splitter "
        "variants is intentionally cautious. Those models are trained under "
        "different data regimes, output representations, and evaluation "
        "protocols [8]. Our contribution is complementary: atom-level "
        "segmentation maps, reconstruction inside the training loss, and an "
        "explicit multi-split OOD benchmark on a transparent CPU-scale "
        "pipeline.",
    )

    add_heading_custom(doc, "6.2 Role of soft reconstruction", level=2)
    add_para(
        doc,
        "Treating reassembly only as an evaluation metric leaves a gap between "
        "what is optimized and what is scientifically desired. The soft "
        "reconstruction loss closes that gap with differentiable proxies that "
        "are cheap to compute yet chemically motivated. The large improvement "
        "on unseen-E3 reassembly after introducing the full reconstruction "
        "term is the strongest empirical argument for retaining reconstruction "
        "inside the methodology rather than only in the results section.",
    )

    add_heading_custom(doc, "6.3 Scratch versus pretrained encoders", level=2)
    add_para(
        doc,
        "AttrMask pretraining does not dominate every split. Scratch training "
        "is preferable on random and unseen-E3 reassembly in our runs, whereas "
        "pretraining helps unseen-warhead and newer-chemotype reassembly. "
        "Reporting both outcomes is scientifically preferable to selecting only "
        "the winning variant. The comparison also satisfies the practical "
        "review expectation that modern molecular learning studies examine "
        "pretrained representations [13,14], even when a compact AttrMask "
        "protocol is used in place of heavier external checkpoints.",
    )

    add_heading_custom(doc, "6.4 Limitations", level=2)
    add_para(
        doc,
        "Several limitations should be stated clearly. Weak labels remain "
        "imperfect; dictionary circularity implies that dictionary scores are "
        "not a fair upper-bound competition. The soft-reconstruction loss is a "
        "proxy rather than an exact differentiable reassembly oracle. AttrMask "
        "on the task backbone is not identical to loading a large public "
        "checkpoint such as GROVER or GraphMVP [17,18]. Finally, absolute "
        "performance on fingerprint-OOD reassembly indicates residual "
        "generalization headroom, motivating richer E3 coverage, chemist-curated "
        "gold subsets, and stronger pretrained encoders in future work.",
    )

    # =====================================================================
    # 7. CONCLUSION
    # =====================================================================
    add_heading_custom(doc, "7. Conclusion", level=1)
    add_para(
        doc,
        "We presented a complete research methodology for atom-level PROTAC "
        "substructure segmentation based on reassembly-filtered weak "
        "supervision, a compact GIN architecture, reconstruction-aware "
        "multi-task training, AttrMask pretraining comparison, and a five-way "
        "OOD evaluation suite. The approach outperforms an XGBoost "
        "bond-cutting baseline on atom accuracy across all studied splits and "
        "shows that soft reconstruction is a high-value addition to the "
        "training objective. The framework is intentionally simple, "
        "reproducible, and suitable for laboratory-scale computation, while "
        "remaining aligned with the scientific requirements of PROTAC "
        "informatics: chemically coherent fragments, honest generalization "
        "tests, and transparent baselines.",
    )

    # =====================================================================
    # 8. FUTURE WORK
    # =====================================================================
    add_heading_custom(doc, "8. Future Work", level=1)
    add_para(
        doc,
        "Promising extensions include curation of a gold-standard atom-labeled "
        "test panel by medicinal chemists; integration of larger pretrained "
        "encoders (Mole-BERT, GROVER, GraphMVP) under controlled fine-tuning; "
        "replacement of soft reconstruction proxies with straight-through or "
        "reinforcement objectives on hard cuts; and external validation on "
        "PROTAC collections beyond PROTAC-DB. Such steps would further "
        "strengthen claims of real-world generalization.",
    )

    # =====================================================================
    # ACKNOWLEDGMENTS
    # =====================================================================
    add_heading_custom(doc, "Acknowledgments", level=1)
    add_para(
        doc,
        "The authors thank contributors of PROTAC-DB and the open-source "
        "scientific Python ecosystem (RDKit, PyTorch, scikit-learn, XGBoost) "
        "that made this study possible. Any opinions expressed are those of "
        "the authors.",
        first_line=False,
    )

    # =====================================================================
    # DATA AND CODE AVAILABILITY
    # =====================================================================
    add_heading_custom(doc, "Data and Code Availability", level=1)
    add_para(
        doc,
        "Processed labels, splits, trained model checkpoints, and evaluation "
        "tables are stored under the project outputs directory. Source code is "
        "organized as a single reusable module with five executable notebooks "
        "covering data preparation, baselines, training, evaluation, and "
        "inference demonstration. PROTAC-DB files should be obtained according "
        "to the database providers’ terms of use [7].",
        first_line=False,
    )

    # =====================================================================
    # REFERENCES
    # =====================================================================
    add_heading_custom(doc, "References", level=1)

    refs = [
        "[1] Békés, M., Langley, D. R., & Crews, C. M. (2022). PROTAC targeted protein degraders: the past is prologue. Nature Reviews Drug Discovery, 21, 181–200.",
        "[2] Chamberlain, P. P., & Hamann, L. G. (2019). Development of targeted protein degradation therapeutics. Nature Chemical Biology, 15, 937–944.",
        "[3] Sakamoto, K. M., et al. (2001). Protacs: chimeric molecules that target proteins to the Skp1–Cullin–F box complex for ubiquitination and degradation. Proceedings of the National Academy of Sciences, 98(15), 8554–8559.",
        "[4] Bondeson, D. P., et al. (2015). Catalytic in vivo protein knockdown by small-molecule PROTACs. Nature Chemical Biology, 11, 611–617.",
        "[5] Troup, R. I., Fallan, C., & Baud, M. G. J. (2020). Current strategies for the design of PROTAC linkers. Exploration of Targeted Anti-tumor Therapy, 1, 273–312.",
        "[6] Cyrus, K., et al. (2011). Impact of linker length on the activity of PROTACs. Molecular BioSystems, 7(2), 359–364.",
        "[7] Weng, G., et al. (2021/updated releases). PROTAC-DB: an online database of PROTACs. Nucleic Acids Research. PROTAC-DB 4.0 resource used in this study.",
        "[8] Ribes, S., et al. (2026). PROTAC-Splitter: a machine learning framework for automated identification of PROTAC substructures. (Published framework combining dictionary curation and learned splitting models).",
        "[9] Sharma, A. K. (2026). Machine learning-driven PROTAC linker identification and related computational analysis. Laboratory internship / project report.",
        "[10] Duvenaud, D., et al. (2015). Convolutional networks on graphs for learning molecular fingerprints. Advances in Neural Information Processing Systems.",
        "[11] Gilmer, J., et al. (2017). Neural message passing for quantum chemistry. Proceedings of Machine Learning Research (ICML).",
        "[12] Xu, K., Hu, W., Leskovec, J., & Jegelka, S. (2019). How powerful are graph neural networks? International Conference on Learning Representations (ICLR).",
        "[13] Hu, W., et al. (2020). Strategies for pre-training graph neural networks. International Conference on Learning Representations (ICLR).",
        "[14] Xia, J., et al. (2023). Mole-BERT: Rethinking pre-training graph neural networks for molecules. International Conference on Learning Representations (ICLR).",
        "[15] Landrum, G., et al. RDKit: Open-source cheminformatics. https://www.rdkit.org",
        "[16] Paszke, A., et al. (2019). PyTorch: An imperative style, high-performance deep learning library. Advances in Neural Information Processing Systems.",
        "[17] Rong, Y., et al. (2020). Self-supervised graph transformer on large-scale molecular data (GROVER). Advances in Neural Information Processing Systems.",
        "[18] Liu, S., et al. (2022). Pre-training molecular graph representation with 3D geometry (GraphMVP). International Conference on Learning Representations (ICLR).",
        "[19] Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining.",
        "[20] Rogers, D., & Hahn, M. (2010). Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 50(5), 742–754.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        pf.space_before = Pt(0)
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0.75)
        pf.first_line_indent = Cm(-0.75)  # hanging indent
        run = p.add_run(r)
        set_run_font(run, size=11)

    # Appendix note
    doc.add_page_break()
    add_heading_custom(doc, "Appendix A. Methodological Schematic", level=1)
    add_para(
        doc,
        "The end-to-end workflow comprises: (i) PROTAC-DB ingestion and SMILES "
        "cleaning; (ii) dictionary matching with reassembly filtering; "
        "(iii) construction of five evaluation splits; (iv) optional AttrMask "
        "pretraining; (v) multi-task GNN fine-tuning with soft reconstruction; "
        "(vi) constrained decoding; and (vii) comparative evaluation against "
        "dictionary and XGBoost baselines. Intermediate artifacts "
        "(labeled_protacs.jsonl, splits.pkl, graphs.pt, gnn_models.pt, "
        "final_results.json) support full reproducibility of the reported "
        "tables and figures.",
        first_line=False,
    )

    add_heading_custom(doc, "Appendix B. Ethical and Reproducibility Statement", level=1)
    add_para(
        doc,
        "This study is computational and does not involve human subjects or "
        "animal experiments. All quantitative claims in the Results section "
        "refer to the metrics stored in the project evaluation exports. The "
        "manuscript text was written to describe the implemented methodology "
        "and observed outcomes in original academic prose. Authors should "
        "replace placeholder affiliation fields and verify citation details "
        "against the final published versions of external papers before "
        "journal submission.",
        first_line=False,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")
    return OUT


if __name__ == "__main__":
    build()
